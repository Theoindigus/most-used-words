# Copyrigth © 2023 Егоров Антон (Egorov Anton) (Theoindigus) <indigus@aegor.ru>
#
# Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation 
# files (the "Software"), to deal in the Software without restriction, including without limitation the rights to use, copy, 
# modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and to permit persons to whom the Soft-
# ware is furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRAN-
# TIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT 
# HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING 
# FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.

from sys import path as abs_path
from sys import platform
from sys import exit
import re
import os.path
from progress.bar import IncrementalBar as IncBar
from docx import Document as Doc

# Подготовка программы
f_not_exists = True # file no exists
f_name = '' # file name
f_path = '' # file path
dir_sep = '' # directory separator
if platform == "linux" or platform == "linux2":
	dir_sep = '/'
elif platform == "win32":
	dir_sep = '\\'
elif platform == "darwin":
	dir_sep = ':'
else:
	print('Unknown platform')
	exit(-1)

# Ввод имени файла и проверка его наличия
while f_not_exists:
	print('Enter the file name. The file should be in the same directory with the script')
	f_name = input()
	f_path = abs_path[0] + dir_sep + f_name + '.docx'
	if not os.path.exists(f_path):
		print("There is no such file: " + f_path)
	else:
		f_not_exists = False

# Создание паттерна по которому из текста будут выделяться слова вместе со знаками пунктуации
# text_pattern = r'\s*([-,.!?:;\'"/\\0-9A-Za-zА-Яа-яЁё]+)\s*'
# Создание паттерна по которому из текста будут выделяться слова без знаков пунктуации
# letters_pattern = r'\s*([-\'"A-Za-zА-Яа-яЁё]+)\s*'

patterns = (r'\s*([-,.!?:;\'"/\\0-9A-Za-z]+)\s*',
			r'\s*([-,.!?:;\'"/\\0-9А-Яа-яЁё]+)\s*',
			r'\s*([-\'"A-Za-z]+)\s*',
			r'\s*([-\'"А-Яа-яЁё]+)\s*')

# Для облегчения обработки файла стоит исключить из него базовый набор
# общеупотребимых слов, а также отдельные буквы алфавита 
excluded_words = set()
if not os.path.exists(abs_path[0] + dir_sep + 'exclude_words.docx'):
	print('The file with the list of excluded words is missing in the script directory')
else:
	for p in Doc(abs_path[0] + dir_sep +'exclude_words.docx').paragraphs:
		excluded_words.add(p.text)


# Получение исходного документа
# Извлечение из него параграфов текста
# Создание прогресс-бара
origin_doc = Doc(f_path)
paragraphs = origin_doc.paragraphs
bar = IncBar('Paragraphs processing:         ', max = len(paragraphs))

# Создание словаря для хранения всех слов без повторов в качестве ключа
# и количества их повторов в качестве значения
# Создание документа, в который будет записан отделённый от всего документа
# текст без форматирования, то есть чистый текст
all_words_set = dict()
plain_text_doc = Doc()

for p in paragraphs:
	match = re.findall(patterns[0], p.text)
	if match:
		plain_text = '\t'
		for m in match:
			if m:
				plain_text += m + ' '
				tmp = re.match(patterns[2], m)
				if tmp:
					tmp = tmp.group(0)
					if tmp not in excluded_words:
						if tmp not in all_words_set:
							all_words_set[tmp.lower()] = 1
						else:
							all_words_set[tmp.lower()] = all_words_set[tmp.lower()] + 1
		plain_text_doc.add_paragraph(plain_text)
	bar.next()
bar.finish()

# Ввод максимального количества наиболее используемых слов, которые надо найти
# Поиск в полученном раннее словаре слов, которые чаще всего встречались в тексте
# Сортировка полученного списка слов по алфавиту
max_search_words = int(input('Enter the maximum number of most used words to be found:\n'))
if max_search_words > len(all_words_set):
	max_search_words = len(all_words_set)
bar = IncBar('Search for the most used words:', max = max_search_words)
most_used_words = list()
for i in range(0, max_search_words):
	max_used_count = 0
	max_used_word = ''
	for e in all_words_set:
		if all_words_set[e] > max_used_count:
			max_used_word = e
			max_used_count = all_words_set[e]
	most_used_words.append(max_used_word)
	del all_words_set[max_used_word]
	bar.next()
most_used_words = sorted(most_used_words)
bar.finish()

# Сохранение в файлы чистого текста и списка слов
bar = IncBar('Files saving:                  ', max = len(most_used_words))
plain_text_doc.save(abs_path[0] + dir_sep + f_name + '_plain_text.docx')
bar.next()
most_used_words_doc = Doc()
for w in most_used_words:
	most_used_words_doc.add_paragraph(w)
	bar.next()
most_used_words_doc.save(abs_path[0] + dir_sep + f_name + '_most_used_words.docx')
bar.finish()
